"""
Helper utilities to load documents from the data directory for ingestion.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List
import csv
import io

from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd


@dataclass
class LoadedDocument:
    title: str
    content: str
    source: str
    metadata: Dict[str, str]


def _strip_gutenberg_header(text: str) -> str:
    """Remove common Project Gutenberg license header/footer if present."""
    lower = text.lower()

    start_idx = 0
    start_marker = "*** start of"
    marker_pos = lower.find(start_marker)
    if marker_pos != -1:
        newline_pos = text.find("\n", marker_pos)
        start_idx = newline_pos + 1 if newline_pos != -1 else marker_pos

    end_idx = len(text)
    end_marker = "*** end of"
    marker_pos = lower.find(end_marker)
    if marker_pos != -1:
        end_idx = marker_pos

    cleaned = text[start_idx:end_idx].lstrip()

    # Remove additional catalog metadata lines (Title, Author, etc.)
    lines = cleaned.splitlines()
    filtered: List[str] = []
    skipping = True
    gutter_prefixes = (
        "the project gutenberg",
        "project gutenberg",
        "title:",
        "author:",
        "editor:",
        "release date:",
        "language:",
        "character set encoding:",
        "produced by",
        "etext prepared by",
        "credits:",
        "illustrator:",
    )
    for line in lines:
        stripped = line.strip()
        if skipping:
            if not stripped:
                continue
            lower_line = stripped.lower()
            if any(lower_line.startswith(prefix) for prefix in gutter_prefixes):
                continue
            skipping = False
        filtered.append(line)

    cleaned = "\n".join(filtered)
    return cleaned


def load_document_from_path(path: str) -> List[LoadedDocument]:
    """Load a text or PDF document and return metadata suitable for ingestion."""
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    if file_path.suffix.lower() == ".txt":
        raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
        text = _strip_gutenberg_header(raw_text)
        return [
            LoadedDocument(
                title=file_path.stem,
                content=text,
                source=file_path.name,
                metadata={"filename": file_path.name},
            )
        ]

    if file_path.suffix.lower() == ".pdf":
        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")

        text = "\n".join(pages)
        text = _strip_gutenberg_header(text)
        return [
            LoadedDocument(
                title=file_path.stem,
                content=text,
                source=file_path.name,
                metadata={
                    "filename": file_path.name,
                    "pages": str(len(reader.pages)),
                },
            )
        ]

    if file_path.suffix.lower() == ".docx":
        doc = DocxDocument(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)
        return [
            LoadedDocument(
                title=file_path.stem,
                content=text,
                source=file_path.name,
                metadata={"filename": file_path.name, "doc_type": "docx"},
            )
        ]

    if file_path.suffix.lower() in {".csv"}:
        with file_path.open("r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)
        # Convert to a simple tab-separated text for ingestion
        text_rows = ["\t".join(row) for row in rows]
        text = "\n".join(text_rows)
        return [
            LoadedDocument(
                title=file_path.stem,
                content=text,
                source=file_path.name,
                metadata={
                    "filename": file_path.name,
                    "doc_type": "csv",
                    "rows": str(len(rows)),
                },
            )
        ]

    if file_path.suffix.lower() in {".xls", ".xlsx"}:
        # Load all sheets and concatenate as text
        sheets = pd.read_excel(file_path, sheet_name=None)
        parts: List[str] = []
        for sheet_name, df in sheets.items():
            buffer = io.StringIO()
            df.to_csv(buffer, index=False)
            parts.append(f"Sheet: {sheet_name}\n{buffer.getvalue()}")

        text = "\n\n".join(parts)
        return [
            LoadedDocument(
                title=file_path.stem,
                content=text,
                source=file_path.name,
                metadata={
                    "filename": file_path.name,
                    "doc_type": "excel",
                    "sheets": str(len(sheets)),
                    "sheet_names": ", ".join(list(sheets.keys())[:5]),
                },
            )
        ]

    raise ValueError(f"Unsupported document type: {path}")
